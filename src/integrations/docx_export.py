"""
docx_export.py — Professional Word Document Export

Turns any council task's final output into a formatted, client-ready
.docx file. Built primarily for the Grant Council so proposal sections
can be downloaded and manually submitted to EU portals (Horizon Europe,
Funding & Tenders, etc.), but works for any council's output.

No external API required — pure python-docx, already a dependency.
"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches


def _safe_output_id(task: dict) -> str:
    """Return a filesystem/header-safe identifier for exported artifacts."""
    task_id = task.get("task_id", "output")
    return re.sub(r"[^a-zA-Z0-9_-]", "", str(task_id)) or "output"

COUNCIL_TITLES = {
    "grant": "Grant Proposal",
    "sales": "Sales Outreach Draft",
    "content": "Content Draft",
}

HEADING_BLUE = RGBColor(0x2E, 0x74, 0xB5)
HEADING_DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x18, 0x18, 0x1B)
MUTED = RGBColor(0x52, 0x52, 0x5B)


def _normalize_export_text(value: str) -> str:
    """Avoid glyph-sensitive dash variants in portal-ready files."""
    return str(value).replace("\u2011", "-").replace("\u2013", "-").replace("\u2014", "-")


def _set_font(style, *, size: float, color: RGBColor, bold: bool = False) -> None:
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold


def _configure_docx(doc: Document) -> None:
    """Apply the grant_proposal preset with an A4 EU-portal page override."""
    section = doc.sections[0]
    section.page_width = Inches(8.2677)
    section.page_height = Inches(11.6929)
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    _set_font(normal, size=11, color=INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, HEADING_BLUE, 16, 8),
        "Heading 2": (13, HEADING_BLUE, 12, 6),
        "Heading 3": (12, HEADING_DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        _set_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        _set_font(style, size=11, color=INK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def _render_body(doc: Document, text: str):
    """
    Render final_output text into the document. Recognizes simple markdown-style
    section markers (### Heading, **bold**, - bullets) so AI-generated structured
    text renders as a real formatted document rather than one flat paragraph.
    """
    lines = _normalize_export_text(text).replace("\r\n", "\n").split("\n")
    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            continue

        heading_match = re.match(r"^#{1,4}\s+(.*)", stripped)
        if heading_match:
            level = min(3, len(stripped) - len(stripped.lstrip("#")))
            doc.add_paragraph(heading_match.group(1), style=f"Heading {level}")
            continue

        bullet_match = re.match(r"^[-*•]\s+(.*)", stripped)
        if bullet_match:
            doc.add_paragraph(bullet_match.group(1), style="List Bullet")
            continue

        numbered_match = re.match(r"^\d+[.)]\s+(.*)", stripped)
        if numbered_match:
            doc.add_paragraph(numbered_match.group(1), style="List Number")
            continue

        p = doc.add_paragraph()
        # Render **bold** inline spans without leaking asterisks into the text.
        segments = re.split(r"(\*\*[^*]+\*\*)", stripped)
        for seg in segments:
            if seg.startswith("**") and seg.endswith("**"):
                run = p.add_run(seg[2:-2])
                run.bold = True
            else:
                run = p.add_run(seg)
            run.font.size = Pt(11)


def build_task_docx(task: dict) -> bytes:
    """
    Build a formatted .docx for a task and return the raw file bytes.
    """
    doc = Document()
    _configure_docx(doc)

    council = (task.get("council") or "").lower()
    title = COUNCIL_TITLES.get(council, "AI Council OS Output")

    doc.core_properties.title = title
    doc.core_properties.author = ""
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    title_p.paragraph_format.keep_with_next = True
    title_run = title_p.add_run(title)
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = INK

    subtitle = task.get("task_description", "")[:200]
    if subtitle:
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.paragraph_format.space_after = Pt(18)
        sr = sp.add_run(subtitle)
        sr.font.name = "Calibri"
        sr.font.size = Pt(12)
        sr.font.color.rgb = MUTED
        sr.italic = True

    final_output = task.get("final_output", "") or "(No final output generated yet.)"
    _render_body(doc, final_output)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_task_docx_filename(task: dict) -> str:
    council = (task.get("council") or "task").lower()
    return f"{council}-{_safe_output_id(task)}.docx"


def build_task_pdf(task: dict) -> bytes:
    """Build a print-ready PDF directly, without requiring LibreOffice."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        BaseDocTemplate,
        Frame,
        ListFlowable,
        ListItem,
        PageTemplate,
        Paragraph,
        Spacer,
    )

    buffer = io.BytesIO()
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "CouncilBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=15,
        textColor=colors.HexColor("#27272A"),
        spaceAfter=6,
        alignment=TA_LEFT,
    )
    heading = ParagraphStyle(
        "CouncilHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#18181B"),
        spaceBefore=8,
        spaceAfter=6,
        keepWithNext=True,
    )
    title_style = ParagraphStyle(
        "CouncilTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=22,
        leading=27,
        textColor=colors.HexColor("#18181B"),
        spaceAfter=8,
        alignment=TA_CENTER,
        keepWithNext=True,
    )
    def _header_footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#71717A"))
        canvas.drawRightString(A4[0] - 20 * mm, 12 * mm, f"Page {doc.page}")
        canvas.restoreState()

    frame = Frame(20 * mm, 18 * mm, A4[0] - 40 * mm, A4[1] - 38 * mm, id="body")
    document = BaseDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=18 * mm,
        title=COUNCIL_TITLES.get((task.get("council") or "").lower(), "AI Council OS Output"),
        author="",
    )
    document.addPageTemplates(PageTemplate(id="main", frames=[frame], onPage=_header_footer))

    council = (task.get("council") or "").lower()
    title = COUNCIL_TITLES.get(council, "AI Council OS Output")
    story = [Paragraph(title, title_style)]
    description = (task.get("task_description") or "").strip()
    if description:
        subtitle_style = ParagraphStyle(
            "CouncilSubtitle",
            parent=body,
            fontSize=11,
            leading=14,
            textColor=colors.HexColor("#52525B"),
            alignment=TA_CENTER,
            spaceAfter=18,
        )
        story.append(Paragraph(_escape_pdf_text(_normalize_export_text(description[:500])), subtitle_style))

    pending_bullets: list[str] = []

    def flush_bullets():
        nonlocal pending_bullets
        if pending_bullets:
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(_inline_pdf_markup(item), body)) for item in pending_bullets],
                    bulletType="bullet",
                    leftIndent=14,
                    bulletFontName="Helvetica",
                    bulletFontSize=8,
                    spaceAfter=6,
                )
            )
            pending_bullets = []

    final_output = _normalize_export_text(task.get("final_output") or "(No final output generated yet.)")
    for raw_line in str(final_output).replace("\r\n", "\n").split("\n"):
        stripped = raw_line.strip()
        if not stripped:
            flush_bullets()
            story.append(Spacer(1, 4))
            continue
        bullet = re.match(r"^[-*•]\s+(.*)", stripped)
        if bullet:
            pending_bullets.append(bullet.group(1))
            continue
        flush_bullets()
        heading_match = re.match(r"^#{1,4}\s+(.*)", stripped)
        if heading_match:
            story.append(Paragraph(_inline_pdf_markup(heading_match.group(1)), heading))
        else:
            story.append(Paragraph(_inline_pdf_markup(stripped), body))
    flush_bullets()

    document.build(story)
    return buffer.getvalue()


def _escape_pdf_text(value: str) -> str:
    from xml.sax.saxutils import escape

    return escape(str(value), {'"': "&quot;", "'": "&apos;"})


def _inline_pdf_markup(value: str) -> str:
    """Escape user/model text while preserving simple **bold** spans."""
    parts = re.split(r"(\*\*[^*]+\*\*)", value)
    rendered: list[str] = []
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            rendered.append(f"<b>{_escape_pdf_text(part[2:-2])}</b>")
        else:
            rendered.append(_escape_pdf_text(part))
    return "".join(rendered)


def build_task_pdf_filename(task: dict) -> str:
    council = (task.get("council") or "task").lower()
    return f"{council}-{_safe_output_id(task)}.pdf"
